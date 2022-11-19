from docx import Document
from datetime import datetime
import os
from research_graph import *
from pykrx import stock
from docx.shared import Inches,Cm



document = Document()
now = datetime.now()
def makeDirectory():
    try:
        if not os.path.isdir("분석 레포트"):
            os.mkdir("분석 레포트")
    except OSError:
        print("err")

def makeReport(code_stock):
    #데이터 불러오기
    now = datetime.now()
    today_date = str(now.date())
    today_date = today_date[0:4] + today_date[5:7] +today_date[8:10]

    perdata = stock.get_market_fundamental_by_date("20150101",today_date,code_stock)
    per_dict = perdata["PER"].to_dict()
    per_list = list(per_dict.values())

    pricedata = stock.get_market_ohlcv_by_date("20150101",today_date,code_stock)

    price_dict = pricedata["종가"].to_dict()
    price_list = list(price_dict.values())
    

    regression_list = draw_all_graph(code_stock)
    #데이터 불러오기 완료


    name_stock = stock.get_market_ticker_name(code_stock)
    price_stock = price_list[-1]
    per_stock = per_list[-1]


    document.add_heading(f'{name_stock} 분석 레포트', level = 0)
    document.add_paragraph(f'본 레포트가 작성된 날짜는{str(now.date())}이며\n 작성된 시간은{str(now.time())}입니다')
    document.add_heading(f"{name_stock}({code_stock})은 현재...")
    document.add_paragraph(f'현재 주가는 {price_stock}원 입니다')
    document.add_paragraph(f'PER 은 {per_stock} 입니다')


    """ 
    # 제목  name_stock,code_stock,price_stock,per_stock,per_other,stock_grade
    # 기본 정보
    document.add_paragraph(f'동종 업계 per은 {per_other} 입니다')
    document.add_paragraph(f'결론적으로, 현재{name_stock}의 투자 매력도는 {stock_grade}% 입니다')
    """

    table = document.add_table(rows=5, cols=2)
    grid_style = document.styles["Table Grid"]
    table.style = grid_style
    # 만든 표의 스타일을 가장 기본 스타일인 'Table Grid'로 설정
    head_cells = table.rows[0].cells
    head_cells[0].text = '그래프'
    head_cells[1].text = '비고'

    for i in range(4):
        cell = table.cell(i+1,0)
        para = cell.add_paragraph()
        run = para.add_run()
        run.add_picture(f"graph {i+1}.png",width = Inches(4),height = Inches(2))
    for i in range(4):
        target_cell = table.rows[i+1].cells
        target_cell[1].text = f"{regression_list[i]}"

    for i in range(4):
        os.remove(f"graph {i+1}.png")

    flist = regression_list[4]
    sum_future = flist[0] * 60 + flist[1] * 30 + flist[2] * 10
    sum_future /= 100
    document.add_paragraph(f'향후 미래 예상 가격 = {sum_future}')
 

    makeDirectory()
    folder_path = os.getcwd() + f'\분석 레포트\ [{str(now.date())}] {name_stock} 분석 레포트.docx'
    document.save(folder_path)

#makeReport(name,code,price,p,po,grade)